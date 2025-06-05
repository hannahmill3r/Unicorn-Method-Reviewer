import docx
from extractText import closest_match_unit_op
import re
from blockNameDict_user_validation import blockNameDictionary

default_process_info = {
        'Regeneration': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre Sanitization Rinse': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre Sanitization Rinse 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Equilibration': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Charge': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre Sanitization': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre Sanitization 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Post Sanitization': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Post Sanitization 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Wash 1': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Wash 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Wash 3': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Elution': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Storage Rinse': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Post Sanitization Rinse': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Post Sanitization Rinse 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Elution': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Storage': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '}, 
        'Neutralization': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '}, 

    }


def read_docx2(file_path):
    """Read content from a docx file separating into nested lists by unit operation"""
    doc = docx.Document(file_path)
    pages = {}
    current_page = []
    first_page = []
    is_first_page = True
    
    def is_unit_op_header(text):
        return "unit" in text.lower() and 'op' in text.lower() and "process flow chart" in text.lower()
    
    def is_downstream_header(text):
        return "Downstream Process Flow Chart" in text
    
    # Process elements in order of appearance
    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('p'):
            paragraph = element.text.strip()
            if paragraph:
                if is_downstream_header(paragraph):
                    current_page.append(paragraph)
                elif is_unit_op_header(paragraph):
                    # Save previous content
                    if is_first_page:
                        pages["Downstream Process Flow Chart"] = current_page
                        is_first_page = False
                    elif current_page:
                        pages[current_page[0]] = current_page
                    # Start new page with Unit Op header
                    current_page = [paragraph]
                else:
                    current_page.append(paragraph)
        
        # Handle tables
        elif element.tag.endswith('tbl'):
            table = doc.tables[len([e for e in doc.element.body[:doc.element.body.index(element)] 
                                  if e.tag.endswith('tbl')])]
            table_content = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_content.append(' | '.join(row_text))
            current_page.extend(table_content)
    
    # Add final page
    if current_page:
        if is_first_page:
            pages["Downstream Process Flow Chart"] = current_page
        else:
            pages[current_page[0]] = current_page

    return pages


def list_unit_ops(pages):
    unitOperations = []
    for i in pages.keys():
        if "Unit Op" in i:
            unitOperations.append(i)

    return unitOperations
    
def extract_process_info(array, unitOP):
    highSaltWash = False
    process_info = {
        'Regeneration': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre Sanitization Rinse': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre Sanitization Rinse 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Equilibration': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Charge': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre Sanitization': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre Sanitization 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Post Sanitization': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Post Sanitization 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Wash 1': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Wash 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Wash 3': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Elution': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Storage Rinse': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Post Sanitization Rinse': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Post Sanitization Rinse 2': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Elution': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Storage': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '}, 
        'Neutralization': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'High Salt Wash': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
        'Pre-Equilibration': {'direction': '', 'velocity': '', 'composition': '','residenceTime': '--', 'CV': ' '},
    }

    parameters_in_pfc = []
    
    default_flow = ''
    current_step = None
    newStep = None
    doubleBuffer = False

    suRe = re.search(r'sure\s*(.*)', array[0].lower())
    sanitizationStrategy = "SuRe" if suRe else "PrismA"

    for item in array:
    
        if parse_default_flow_direction(item.lower()):
            default_flow = parse_default_flow_direction(item.lower())
        currentHeader = False
        parts = item.split('|')
        parts = [p.strip() for p in parts]
        lower_item = item.lower()
        if "/" in lower_item:
            current_steps = []
            lower_items = lower_item.split('/')
            for i in lower_items:
                if "step" in i:
                    currentHeader = True
            for itemInBackSlash in lower_items:
                step = detect_PFC_step(itemInBackSlash, currentHeader, unitOP)
                if step and step not in current_steps:
                    current_steps.append(step)

            if len(current_steps)>=2:
                current_step = ', '.join(current_steps)
                doubleBuffer = True
            else:
                doubleBuffer = False
        if not doubleBuffer or "/" not in lower_item:
            if detect_PFC_step(lower_item, currentHeader, unitOP):
                current_step = detect_PFC_step(lower_item, currentHeader, unitOP)

        # If we're in a process step, capture its parameters
        if current_step and 'step' not in lower_item:
            print(current_step, lower_item)
            newStep = current_step.split(', ')
            if 'rinse' in lower_item:
                if current_step + " Rinse" in process_info.keys():
                    newStep = [current_step + " Rinse"]
                
            # Flow direction
            if 'flow direction' in lower_item:
                for buffer in newStep:
                    process_info[buffer]['direction'] = parts[1] if len(parts) > 1 else ''
            # Flow velocity - only capturing NMT value
            elif 'cv' in parts[0].lower() and 'volume' in  parts[0].lower():
                for buffer in newStep:
                    process_info[buffer]['CV'] = parts[1] if len(parts) > 1 else ''
            elif any(term in lower_item for term in ['flow velocity', 'velocity', 'flow:', 'residence', 'exposure']):

                # Extract the NMT value
                try:
                    value = item[item.lower().find('nmt'):].split()[1]

                    for buffer in newStep:
                        process_info[buffer]['velocity'] = value

                except:
                    pass

                try:
                    
                    value = item[item.lower().find('nlt'):].split()[1]
                    for buffer in newStep:
                        process_info[buffer]['residenceTime'] = value

                except:
                    pass
                salt_indicators = ['nacl', 'kcl', 'cacl', 'mgcl', 'na2so4', 'K2SO4',
                                'sodium chloride', 'potassium chloride', 'calcium chloride',
                                'sodium acetate', 'potassium acetate',
                                'sodium phosphate', 'potassium phosphate'
                            ]
            # Composition
            elif any(term in lower_item for term in ['composition', 'buffer composition']):
                for buffer in newStep:
                    #In Non ProA Chromatography, high salt washes may replace regeneration steps if the composition is a salt
                    if buffer == "Regeneration":
                        if unitOP!= "Protein A Capture Chromatography":
                            if any(term in lower_item for term in salt_indicators):
                                highSaltWash = True

                        


                    process_info[buffer]['composition'] = parts[1] if len(parts) > 1 else ''
    print(process_info)
    process_info.pop('Neutralization')
    if highSaltWash:
        process_info['High Salt Wash'] = process_info['Regeneration']
        process_info.pop('Regeneration')

    #grab all of the buffers that had something explicitly included about them in the PFC, will ignore all of the other ones later when we display them with streamlit
    for buffer in process_info.keys():
        for key in process_info[buffer].keys():
            
            if process_info[buffer][key].strip()!='' and process_info[buffer][key].strip()!='--'and buffer not in parameters_in_pfc:
                parameters_in_pfc.append(buffer)

    for key in process_info['Pre-Equilibration'].keys():
        if key!= 'CV':
            if process_info['Pre-Equilibration'][key].strip() == '' and process_info['Equilibration'][key].strip() != '':
                process_info['Pre-Equilibration'][key] = process_info['Equilibration'][key]

    for key in process_info['Wash 1'].keys():
        if key!= 'CV':
            if process_info['Wash 2'][key].strip() == '' and process_info['Wash 1'][key].strip() != '':
                process_info['Wash 2'][key] = process_info['Wash 1'][key]
            if process_info['Wash 3'][key].strip() == '' and process_info['Wash 1'][key].strip() != '':
                process_info['Wash 3'][key] = process_info['Wash 1'][key]


    #parameters are assumed to be shared across a rinse and a buffer step if it is not specified in the pfc for proA
    if unitOP == "Protein A Capture Chromatography":
        checks = ['direction', 'velocity', 'composition']
        for key in process_info.keys():
            if key in parameters_in_pfc:
                for param in checks:
                    if key + " Rinse" in process_info.keys() and process_info[key][param].strip() !='' and process_info[key + ' Rinse'][param].strip() == '':
                        process_info[key + ' Rinse'][param] = process_info[key][param].strip()

    # Set default flow direction if not specified
    for key in process_info:
        if key in parameters_in_pfc:
            if process_info[key]['direction'] == '' and default_flow:
                process_info[key]['direction'] = default_flow

    return process_info, sanitizationStrategy, parameters_in_pfc



def output_PFC_params(PFCInput, unitOperationInMethod):

    textPages = read_docx2(PFCInput)

    unit_operations = list_unit_ops(textPages)

    bestMatchUnitOp, ratio = closest_match_unit_op(unitOperationInMethod, unit_operations)
        
    for unitOp in unit_operations:
        if bestMatchUnitOp.lower() in unitOp.lower():
            details = textPages.get(unitOp)
            process_info, saniStrategy, parameters_in_pfc = extract_process_info(details, unitOperationInMethod)

            return process_info, saniStrategy, parameters_in_pfc
    return default_process_info, "None", []




def parse_default_flow_direction(text):
    # Convert text to lowercase and standardize spacing
    text = ' '.join(text.lower().split())
    
    # Define patterns to match various forms of the statement
    patterns = [
        # Pattern for "all column directions/flows are X unless otherwise noted"
        r'all\s+column\s*(?:directions?|flows?)\s+(?:are|is)\s+(up|down)(?:flow|[-\s]?flow)\s+unless\s+otherwise\s+noted',
        
        # Pattern for "column direction/flow is X unless otherwise noted"
        r'column\s*(?:directions?|flows?)\s+(?:are|is)\s+(up|down)(?:flow|[-\s]?flow)\s+unless\s+otherwise\s+noted',
        
        # Pattern for "X flow unless otherwise noted"
        r'(up|down)(?:flow|[-\s]?flow)\s+unless\s+otherwise\s+noted'
    ]
    
    # Try each pattern
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            direction = match.group(1).lower()
            if "up" in direction.lower():
                return "Upflow"
            elif "down" in direction.lower():
                return "Downflow"
            else:
                return None
    
    return None


def detect_PFC_step(lower_item, currentHeader, unitOp):
    # Detect current process step
    lower_item = lower_item.split('|')[0]
    current_step = None
   
    if 'elution' in lower_item:
        current_step = 'Elution'
    elif 'column preparation' in lower_item:
        if unitOp== "Protein A Capture Chromatography": 
            current_step = 'Equilibration, Pre Sanitization Rinse, Pre Sanitization' 
        else:
            current_step = 'Equilibration, Pre Sanitization'
    elif 'neutralization' in lower_item:
        current_step = 'Neutralization'
    elif 'charge' in lower_item:
        current_step = 'Charge'
    elif 'wash' in lower_item:
        if 'wash 2' in lower_item:
            current_step = 'Wash 2'
        elif 'wash 3' in lower_item:
            current_step = 'Wash 3'
        else:
            current_step = 'Wash 1'

    elif'regeneration' in lower_item:
        current_step = 'Regeneration'
    elif 'pre' in lower_item and 'equil' in lower_item:
        print("test")
        current_step = 'Pre-Equilibration'
    elif 'equil' in lower_item:
        current_step = 'Equilibration'
    elif 'pre' in lower_item and ('use' in lower_item or 'sani' in lower_item) and 'rinse' in lower_item:
        current_step = 'Pre Sanitization Rinse'
    elif 'pre' in lower_item and ('use' in lower_item or 'sani' in lower_item) and 'rinse' not in lower_item:
        current_step = 'Pre Sanitization'
    elif 'sanitization' in lower_item:
        current_step = 'Post Sanitization'
    elif 'storage' in lower_item:
        current_step = 'Storage'

    return current_step
